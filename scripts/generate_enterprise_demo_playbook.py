"""Generate Enterprise Demo Playbook (Word + Interactive HTML)."""
from __future__ import annotations

import html
import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from demo_playbook_content import (
    BASE_URL,
    BLOCKS,
    COMMERCIAL_TIPS,
    COMPANY,
    CREDENTIALS,
    DURING_DEMO_CHECKLIST,
    GLOBAL_FAQ,
    JOURNEY_NARRATIVE,
    POST_DEMO_CHECKLIST,
    PRE_DEMO_CHECKLIST,
    ROLES_OVERVIEW,
    TENANT_ID,
    enrich_blocks,
)

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "demo"
DOCX_OUT = OUT_DIR / "COMPLIANCE360_ENTERPRISE_DEMO_PLAYBOOK.docx"
HTML_OUT = OUT_DIR / "COMPLIANCE360_INTERACTIVE_DEMO_GUIDE.html"
VISUAL_HTML_OUT = OUT_DIR / "COMPLIANCE360_VISUAL_DEMO_GUIDE.html"


def add_para(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("COMPLIANCE 360")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Enterprise Customer Demo Playbook")
    sr.bold = True
    sr.font.size = Pt(18)

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2.add_run(
        f"Customer Journey End-to-End\n"
        f"Escenario: {COMPANY['name']}\n"
        f"{datetime.now(timezone.utc).strftime('%Y-%m-%d')}\n"
        "Confidential — Commercial Use Only"
    )
    doc.add_page_break()

    # TOC placeholder
    doc.add_heading("Índice", 1)
    for i, b in enumerate(BLOCKS, 1):
        add_para(doc, f"{i}. {b['title']} ({b['duration']})")
    add_para(doc, "A. Roles y credenciales")
    add_para(doc, "B. Checklists")
    add_para(doc, "C. FAQ global")
    add_para(doc, "D. Recomendaciones comerciales")
    doc.add_page_break()

    # Scenario
    doc.add_heading("Escenario de demostración", 1)
    add_para(doc, f"Empresa: {COMPANY['name']}", bold=True)
    add_para(doc, f"Industria: {COMPANY['industry']}")
    add_para(doc, f"Normativas: {COMPANY['standards']}")
    add_para(doc, "Problemas actuales del cliente:", bold=True)
    for p in COMPANY["pain_points"]:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("Secuencia narrativa optimizada", 2)
    add_para(
        doc,
        "Esta demo NO sigue el orden del menú. Sigue una historia de negocio: "
        "dolor → plataforma → tenant → roles → documentos → SoD → auditoría → CAPA → "
        "riesgos → KPIs → reportes → infraestructura → viewer → dashboard → valor → cierre."
    )

    doc.add_heading("Diagrama del ciclo de cumplimiento (Customer Journey)", 2)
    flow = doc.add_table(rows=5, cols=3)
    flow.style = "Table Grid"
    steps = [
        ("1. GOBIERNO", "Platform Admin → Tenant Admin → RBAC", "Bloques 2–5"),
        ("2. DOCUMENTOS", "Elaborar (Doc Controller) → Aprobar (QM)", "Bloques 6–7"),
        ("3. AUDITORÍA", "Programa → Ejecución → Hallazgos", "Bloques 8–9"),
        ("4. MEJORA", "CAPA → Riesgos → Indicadores", "Bloques 10–12"),
        ("5. VALOR", "Reportes → Storage → Dashboard → Cierre", "Bloques 13–20"),
    ]
    for i, (phase, detail, blocks) in enumerate(steps):
        flow.rows[i].cells[0].text = phase
        flow.rows[i].cells[1].text = detail
        flow.rows[i].cells[2].text = blocks

    doc.add_paragraph()
    add_para(doc, "Flujo de información entre módulos:", bold=True)
    info_flow = doc.add_table(rows=1, cols=7)
    info_flow.style = "Table Grid"
    chain = ["Documentos", "→", "Auditorías", "→", "Hallazgos", "→", "CAPA"]
    for i, c in enumerate(chain):
        info_flow.rows[0].cells[i].text = c
    row2 = info_flow.add_row().cells
    row2[0].text = "Riesgos"
    row2[1].text = "↔"
    row2[2].text = "Indicadores"
    row2[3].text = "→"
    row2[4].text = "Reportes"
    row2[5].text = "→"
    row2[6].text = "Dashboard"

    doc.add_page_break()

    # Roles table
    doc.add_heading("Roles en la demostración", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Rol", "Propósito", "Puede", "No puede", "Valor"]):
        hdr[i].text = h
    for row_data in ROLES_OVERVIEW:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_page_break()

    # Blocks
    for block in BLOCKS:
        doc.add_heading(f"Bloque {block['id']}: {block['title']}", 1)
        meta = [
            ("Duración", block["duration"]),
            ("Rol", block["role"]),
            ("Usuario", block["user"]),
            ("Pantalla", block["screen"]),
            ("Ruta", block["route"]),
            ("Menú", block["menu"]),
        ]
        mt = doc.add_table(rows=len(meta), cols=2)
        mt.style = "Table Grid"
        for i, (k, v) in enumerate(meta):
            mt.rows[i].cells[0].text = k
            mt.rows[i].cells[1].text = str(v)

        doc.add_heading("Qué hacer", 2)
        for action in block["actions"]:
            doc.add_paragraph(action, style="List Number")

        add_para(doc, f"Botones: {block['buttons']}", bold=True)
        add_para(doc, f"Datos a ingresar: {block['data']}", bold=True)
        add_para(doc, f"Resultado a mostrar: {block['result']}", bold=True)

        doc.add_heading("Qué decir (Customer Script)", 2)
        add_para(doc, block["script"])

        doc.add_heading("Beneficio para el cliente", 2)
        add_para(doc, block["benefit"])

        if block["faq"]:
            doc.add_heading("Preguntas posibles y respuestas", 2)
            for q, a in block["faq"]:
                add_para(doc, f"P: {q}", bold=True)
                add_para(doc, f"R: {a}")

        add_para(doc, f"Advertencias: {block['warnings']}", bold=True)
        add_para(doc, f"Siguiente paso: Bloque {block['id'] + 1 if block['id'] < 20 else '— Fin'}")
        doc.add_page_break()

    # Credentials
    doc.add_heading("Anexo A — Credenciales de demo", 1)
    ct = doc.add_table(rows=1, cols=4)
    ct.style = "Table Grid"
    ct.rows[0].cells[0].text = "Rol"
    ct.rows[0].cells[1].text = "Email"
    ct.rows[0].cells[2].text = "Password"
    ct.rows[0].cells[3].text = "Tenant ID"
    for role, (email, pwd, tid) in CREDENTIALS.items():
        row = ct.add_row().cells
        row[0].text = role
        row[1].text = email
        row[2].text = pwd
        row[3].text = tid

    # Checklists
    doc.add_heading("Anexo B — Checklist previo a la demo", 1)
    for item in PRE_DEMO_CHECKLIST:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")

    doc.add_heading("Checklist durante la demo", 1)
    for item in DURING_DEMO_CHECKLIST:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")

    doc.add_heading("Checklist posterior a la demo", 1)
    for item in POST_DEMO_CHECKLIST:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")

    doc.add_heading("Anexo C — FAQ global", 1)
    for q, a in GLOBAL_FAQ:
        add_para(doc, f"P: {q}", bold=True)
        add_para(doc, f"R: {a}")

    doc.add_heading("Anexo D — Recomendaciones comerciales", 1)
    for tip in COMMERCIAL_TIPS:
        doc.add_paragraph(tip, style="List Bullet")

    doc.add_heading("Cierre ejecutivo", 1)
    add_para(
        doc,
        "Compliance 360 está certificado funcionalmente (29/29 E2E, 23/23 customer journey, "
        "238 unit tests). Esta demo muestra un producto Enterprise listo para piloto comercial "
        "con configuración third-party pendiente (SMTP, storage cloud, SSO)."
    )

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


def _demo_payload() -> dict:
    return {
        "company": COMPANY,
        "baseUrl": BASE_URL,
        "tenantId": TENANT_ID,
        "credentials": {k: {"email": v[0], "password": v[1], "tenantId": v[2]} for k, v in CREDENTIALS.items()},
        "blocks": [
            {
                **block,
                "faq": [{"q": q, "a": a} for q, a in block.get("faq", [])],
            }
            for block in enrich_blocks(BLOCKS)
        ],
        "rolesOverview": [
            {"role": r[0], "purpose": r[1], "can": r[2], "cannot": r[3], "value": r[4]}
            for r in ROLES_OVERVIEW
        ],
        "globalFaq": [{"q": q, "a": a} for q, a in GLOBAL_FAQ],
        "preDemo": PRE_DEMO_CHECKLIST,
        "duringDemo": DURING_DEMO_CHECKLIST,
        "postDemo": POST_DEMO_CHECKLIST,
        "commercialTips": COMMERCIAL_TIPS,
        "closingScript": BLOCKS[-1]["script"],
        "closingActions": BLOCKS[-1]["actions"],
        "journey": JOURNEY_NARRATIVE,
    }


def build_html() -> None:
    payload = _demo_payload()
    data_json = json.dumps(payload, ensure_ascii=False)
    scripts_dir = Path(__file__).resolve().parent
    shared_js = (scripts_dir / "demo_guide_shared.js").read_text(encoding="utf-8")
    template = (scripts_dir / "demo_guide_template.html").read_text(encoding="utf-8")
    html_content = template.replace("__DATA_JSON__", data_json).replace("__SHARED_JS__", shared_js)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    HTML_OUT.write_text(html_content, encoding="utf-8")
    print(f"Wrote {HTML_OUT}")


def build_visual_html() -> None:
    payload = _demo_payload()
    data_json = json.dumps(payload, ensure_ascii=False)
    scripts_dir = Path(__file__).resolve().parent
    shared_js = (scripts_dir / "demo_guide_shared.js").read_text(encoding="utf-8")
    template = (scripts_dir / "demo_visual_guide_template.html").read_text(encoding="utf-8")
    html_content = template.replace("__DATA_JSON__", data_json).replace("__SHARED_JS__", shared_js)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    VISUAL_HTML_OUT.write_text(html_content, encoding="utf-8")
    print(f"Wrote {VISUAL_HTML_OUT}")


def main() -> None:
    build_docx()
    build_html()
    build_visual_html()


if __name__ == "__main__":
    main()
